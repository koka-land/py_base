# schedule_parser.py
from docx import Document
import json
from datetime import datetime
import re


def parse_schedule_to_structured_list(file_path):
    """
    Преобразует таблицу расписания в структурированный список событий
    Каждое событие - словарь с полями:
    - date: дата (например, "1 декабря")
    - weekday: день недели
    - day_full: полное описание дня (например, "1 декабря понедельник")
    - calendar: календарные пометки (святые, праздники)
    - time: время службы
    - service: тип богослужения
    - priest: священник
    - is_special: является ли особой службой (праздник, полиелей и т.д.)
    """

    print("Обработка документа...")
    doc = Document(file_path)

    if not doc.tables:
        print("В документе нет таблиц")
        return []

    table = doc.tables[0]
    structured_events = []
    current_date_info = {}

    print(f"Обрабатывается таблица из {len(table.rows)} строк...")

    for row_idx, row in enumerate(table.rows):
        # Пропускаем заголовок таблицы
        if row_idx == 0:
            continue

        # Получаем текст из ячеек
        cells = [cell.text.strip() for cell in row.cells]

        # Если ячейка 0 содержит дату, обновляем текущую дату
        if cells and cells[0]:
            # Парсим информацию о дне из первой колонки
            day_info = parse_day_info(cells[0])
            if day_info:
                current_date_info = day_info

        # Если есть время в третьей колонке, создаем событие
        if len(cells) >= 3 and cells[2]:
            event = create_event_from_row(current_date_info, cells)
            if event:
                structured_events.append(event)

    print(f"\nОбработано {len(structured_events)} событий")
    return structured_events


def parse_day_info(day_cell_text):
    """
    Парсит информацию о дне из текста первой колонки
    """
    if not day_cell_text:
        return {}

    # Паттерны для извлечения даты и дня недели
    date_patterns = [
        r'(\d{1,2}\s+[а-яёА-ЯЁ]+)\s+([а-яёА-ЯЁ]+)',  # "1 декабря понедельник"
        r'(\d{1,2}\s+[а-яёА-ЯЁ]+)',  # Просто дата
    ]

    day_info = {}
    day_info['day_full'] = day_cell_text

    for pattern in date_patterns:
        match = re.search(pattern, day_cell_text, re.IGNORECASE)
        if match:
            if len(match.groups()) >= 2:
                day_info['date'] = match.group(1)  # "1 декабря"
                day_info['weekday'] = match.group(2).lower()  # "понедельник"
            else:
                day_info['date'] = match.group(1)

            # Извлекаем номер дня
            day_num = re.search(r'^\d{1,2}', day_cell_text)
            if day_num:
                day_info['day_number'] = int(day_num.group())

            return day_info

    return day_info


def create_event_from_row(date_info, cells):
    """
    Создает структурированное событие из строки таблицы
    """
    event = {}

    # Копируем информацию о дате
    event.update(date_info)

    # Основные поля
    event['calendar'] = clean_text(cells[1]) if len(cells) > 1 else ''
    event['time'] = clean_text(cells[2])
    event['service'] = clean_text(cells[3]) if len(cells) > 3 else ''
    event['priest'] = clean_text(cells[4]) if len(cells) > 4 else ''

    # Определяем тип события
    event['event_type'] = determine_event_type(event['service'], event['calendar'])

    # Флаги для фильтрации
    event['is_liturgy'] = is_liturgy(event['service'])
    event['is_vigil'] = is_vigil(event['service'])
    event['is_moleben'] = is_moleben(event['service'])
    event['is_special'] = is_special_service(event['service'], event['calendar'])

    # Извлекаем информацию о праздниках
    event['holidays'] = extract_holidays(event['calendar'])

    # Дополнительная обработка времени
    event['time_parsed'] = parse_time(event['time'])

    return event


def clean_text(text):
    """Очищает текст от лишних пробелов и форматирования"""
    if not text:
        return ""

    # Заменяем множественные пробелы и переносы строк
    text = re.sub(r'\s+', ' ', text)
    # Убираем лишние символы форматирования таблицы
    text = text.replace('**', '').replace('++', '')
    return text.strip()


def determine_event_type(service_text, calendar_text):
    """Определяет тип события"""
    service_lower = service_text.lower()
    calendar_lower = calendar_text.lower()

    if 'литургия' in service_lower:
        return 'liturgy'
    elif 'всенощн' in service_lower:
        return 'vigil'
    elif 'вечерня' in service_lower or 'утреня' in service_lower:
        return 'vespers_matins'
    elif 'молебен' in service_lower:
        return 'moleben'
    elif 'соборование' in service_lower:
        return 'unction'
    elif 'панихида' in service_lower:
        return 'panikhida'
    else:
        return 'other'


def is_liturgy(service_text):
    return 'литургия' in service_text.lower()


def is_vigil(service_text):
    return 'всенощн' in service_text.lower()


def is_moleben(service_text):
    return 'молебен' in service_text.lower()


def is_special_service(service_text, calendar_text):
    """Определяет, является ли служба особой (праздничной и т.д.)"""
    service_lower = service_text.lower()
    calendar_lower = calendar_text.lower()

    special_keywords = [
        'полиелей', 'великое славословие', 'праздник',
        'введение', 'николай', 'андрея первозванного',
        'водосвятный', 'соборование'
    ]

    for keyword in special_keywords:
        if keyword in service_lower or keyword in calendar_lower:
            return True

    return False


def extract_holidays(calendar_text):
    """Извлекает упоминания о праздниках и святых"""
    if not calendar_text:
        return []

    holidays = []

    # Паттерны для поиска праздников
    patterns = [
        r'«([^»]+)»',  # Названия в кавычках
        r'\*\*([^*]+)\*\*',  # Выделенные жирным
        r'\[([^\]]+)\]',  # В квадратных скобках
    ]

    for pattern in patterns:
        matches = re.findall(pattern, calendar_text)
        holidays.extend(matches)

    # Также ищем ключевые слова
    holiday_keywords = [
        'праздник', 'святитель', 'преподобный', 'мученик',
        'апостол', 'пророк', 'благоверный', 'священномученик'
    ]

    # Находим фразы с ключевыми словами
    words = calendar_text.split()
    for i, word in enumerate(words):
        if word.lower() in holiday_keywords and i + 1 < len(words):
            holiday_phrase = ' '.join(words[max(0, i - 1):min(len(words), i + 4)])
            if holiday_phrase not in holidays:
                holidays.append(holiday_phrase)

    return list(set(holidays))  # Убираем дубликаты


def parse_time(time_text):
    """Парсит время из текста"""
    if not time_text:
        return {}

    # Ищем время в формате HH:MM
    time_match = re.search(r'(\d{1,2}):(\d{2})', time_text)

    if time_match:
        hour = int(time_match.group(1))
        minute = int(time_match.group(2))

        # Определяем часть дня
        if hour < 12:
            time_of_day = 'утро'
        elif hour < 17:
            time_of_day = 'день'
        elif hour < 21:
            time_of_day = 'вечер'
        else:
            time_of_day = 'ночь'

        return {
            'hour': hour,
            'minute': minute,
            'time_of_day': time_of_day,
            'formatted': f"{hour:02d}:{minute:02d}"
        }

    return {}


def save_structured_data(events, output_format='json'):
    """
    Сохраняет структурированные данные в разных форматах
    """
    base_filename = 'schedule_december_2025'

    if output_format == 'json':
        filename = f'{base_filename}.json'
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(events, f, ensure_ascii=False, indent=2, default=str)
        print(f"Данные сохранены в {filename}")

    elif output_format == 'json_minimal':
        # Упрощенная версия JSON
        minimal_events = []
        for event in events:
            minimal = {
                'date': event.get('date', ''),
                'weekday': event.get('weekday', ''),
                'time': event.get('time', ''),
                'service': event.get('service', ''),
                'priest': event.get('priest', ''),
                'event_type': event.get('event_type', '')
            }
            minimal_events.append(minimal)

        filename = f'{base_filename}_minimal.json'
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(minimal_events, f, ensure_ascii=False, indent=2)
        print(f"Упрощенные данные сохранены в {filename}")

    elif output_format == 'grouped_json':
        # Группировка по дням
        grouped_by_day = {}
        for event in events:
            day_key = event.get('date', 'unknown')
            if day_key not in grouped_by_day:
                grouped_by_day[day_key] = {
                    'date': day_key,
                    'weekday': event.get('weekday', ''),
                    'day_full': event.get('day_full', ''),
                    'events': []
                }

            # Добавляем только основные поля события
            event_summary = {
                'time': event.get('time', ''),
                'service': event.get('service', ''),
                'priest': event.get('priest', ''),
                'event_type': event.get('event_type', '')
            }
            grouped_by_day[day_key]['events'].append(event_summary)

        # Преобразуем в список для сохранения
        grouped_list = list(grouped_by_day.values())

        filename = f'{base_filename}_grouped.json'
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(grouped_list, f, ensure_ascii=False, indent=2)
        print(f"Сгруппированные данные сохранены в {filename}")

    return events


def analyze_schedule(events):
    """Анализирует расписание и выводит статистику"""
    print("\n" + "=" * 60)
    print("АНАЛИЗ РАСПИСАНИЯ")
    print("=" * 60)

    # Общая статистика
    print(f"\n📊 ОБЩАЯ СТАТИСТИКА:")
    print(f"   Всего событий: {len(events)}")

    # Статистика по типам событий
    event_types = {}
    for event in events:
        e_type = event.get('event_type', 'other')
        event_types[e_type] = event_types.get(e_type, 0) + 1

    print(f"\n📅 ТИПЫ СОБЫТИЙ:")
    for e_type, count in event_types.items():
        print(f"   {e_type}: {count}")

    # Священники
    priests = {}
    for event in events:
        priest = event.get('priest', 'Не указан')
        if priest:
            priests[priest] = priests.get(priest, 0) + 1

    print(f"\n🙏 СВЯЩЕННИКИ (количество служб):")
    for priest, count in sorted(priests.items(), key=lambda x: x[1], reverse=True):
        print(f"   {priest}: {count}")

    # Временной анализ
    morning_services = sum(1 for e in events if e.get('time_parsed', {}).get('time_of_day') == 'утро')
    evening_services = sum(1 for e in events if e.get('time_parsed', {}).get('time_of_day') == 'вечер')

    print(f"\n⏰ ВРЕМЕННОЙ АНАЛИЗ:")
    print(f"   Утренние службы: {morning_services}")
    print(f"   Вечерние службы: {evening_services}")

    # Особые службы
    special_services = [e for e in events if e.get('is_special')]
    print(f"   Особые службы: {len(special_services)}")

    # Примеры событий
    print(f"\n📝 ПРИМЕРЫ СОБЫТИЙ:")
    for i, event in enumerate(events[:3]):
        print(f"\n   Пример {i + 1}:")
        print(f"     Дата: {event.get('date', '')} ({event.get('weekday', '')})")
        print(f"     Время: {event.get('time', '')}")
        print(f"     Служба: {event.get('service', '')}")
        print(f"     Священник: {event.get('priest', '')}")

    return {
        'total_events': len(events),
        'event_types': event_types,
        'priests': priests,
        'morning_services': morning_services,
        'evening_services': evening_services,
        'special_services': len(special_services)
    }


def filter_events(events, filters=None):
    """Фильтрует события по заданным критериям"""
    if not filters:
        return events

    filtered = events.copy()

    # Применяем фильтры
    if 'priest' in filters:
        priest_name = filters['priest'].lower()
        filtered = [e for e in filtered if priest_name in e.get('priest', '').lower()]

    if 'event_type' in filters:
        filtered = [e for e in filtered if e.get('event_type') == filters['event_type']]

    if 'date' in filters:
        filtered = [e for e in filtered if filters['date'] in e.get('date', '')]

    if 'time_of_day' in filters:
        filtered = [e for e in filtered if
                    e.get('time_parsed', {}).get('time_of_day') == filters['time_of_day']]

    return filtered


# Основной код программы
if __name__ == "__main__":
    try:
        # 1. Парсим документ
        print("🔍 Начинаю обработку документа dec_2025.docx...")
        structured_events = parse_schedule_to_structured_list("schedule/dec_2025.docx")

        if not structured_events:
            print("Не удалось извлечь события")
            exit()

        # 2. Анализируем расписание
        stats = analyze_schedule(structured_events)

        # 3. Сохраняем в разных форматах
        print("\n💾 Сохранение данных...")
        save_structured_data(structured_events, 'json')
        save_structured_data(structured_events, 'json_minimal')
        save_structured_data(structured_events, 'grouped_json')

        # 4. Примеры фильтрации
        print("\n🔎 ПРИМЕРЫ ФИЛЬТРАЦИИ:")

        # Все службы прот. Алексея
        print("\n1. Все службы протоиерея Алексея:")
        alexei_events = filter_events(structured_events, {'priest': 'Алексей'})
        for event in alexei_events[:5]:  # Показываем первые 5
            print(f"   {event.get('date')} {event.get('time')}: {event.get('service')}")

        # Все литургии
        print("\n2. Все Божественные Литургии:")
        liturgies = filter_events(structured_events, {'event_type': 'liturgy'})
        print(f"   Найдено {len(liturgies)} литургий")

        # Службы 4 декабря
        print("\n3. Расписание на 4 декабря:")
        dec4_events = filter_events(structured_events, {'date': '4 декабря'})
        for event in dec4_events:
            print(f"   {event.get('time')}: {event.get('service')} - {event.get('priest')}")

        # 5. Выводим общую информацию
        print("\n" + "=" * 60)
        print("✅ ОБРАБОТКА ЗАВЕРШЕНА")
        print("=" * 60)
        print(f"\nСозданы файлы:")
        print("  • schedule_december_2025.json - полные структурированные данные")
        print("  • schedule_december_2025_minimal.json - упрощенный формат")
        print("  • schedule_december_2025_grouped.json - сгруппировано по дням")

        # 6. Пример доступа к данным
        print(f"\nПример работы с данными в Python:")
        print(f"  Всего событий в списке: {len(structured_events)}")
        print(f"  Первое событие: {structured_events[0].get('date')} в {structured_events[0].get('time')}")
        print(f"  Тип первого события: {structured_events[0].get('event_type')}")

    except FileNotFoundError:
        print("❌ Ошибка: Файл dec_2025.docx не найден!")
        print("Убедитесь, что файл находится в той же папке, что и скрипт.")
    except Exception as e:
        print(f"❌ Произошла ошибка: {e}")
        import traceback

        traceback.print_exc()